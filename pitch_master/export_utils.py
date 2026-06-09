"""Pitch Master — Export Utilities."""

from __future__ import annotations

import os
import datetime
from pitch_master.config import OUTPUT_DIR


def _timestamp() -> str:
    return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")


def export_markdown(content: str, prefix: str = "pitch") -> str:
    """Export content as Markdown file. Returns file path."""
    filename = f"{prefix}_{_timestamp()}.md"
    filepath = os.path.join(OUTPUT_DIR, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
    return filepath


def export_txt(content: str, prefix: str = "pitch") -> str:
    """Export content as plain text file. Returns file path."""
    filename = f"{prefix}_{_timestamp()}.txt"
    filepath = os.path.join(OUTPUT_DIR, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
    return filepath


def export_docx(content: str, prefix: str = "pitch") -> str:
    """Export content as DOCX file. Returns file path."""
    from docx import Document

    filename = f"{prefix}_{_timestamp()}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)

    doc = Document()
    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("---"):
            doc.add_paragraph("")
        elif stripped.startswith("**") and stripped.endswith("**"):
            doc.add_paragraph(stripped[2:-2], style="List Bullet")
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)

    doc.save(filepath)
    return filepath


def export_pdf(content: str, prefix: str = "pitch") -> str:
    """Export content as PDF file. Returns file path."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch

    filename = f"{prefix}_{_timestamp()}.pdf"
    filepath = os.path.join(OUTPUT_DIR, filename)

    doc = SimpleDocTemplate(filepath, pagesize=A4,
                            topMargin=0.5*inch, bottomMargin=0.5*inch,
                            leftMargin=0.75*inch, rightMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle("TitleCustom", parent=styles["Title"], fontSize=18, spaceAfter=12)
    heading_style = ParagraphStyle("HeadingCustom", parent=styles["Heading2"], fontSize=14, spaceAfter=8)
    body_style = ParagraphStyle("BodyCustom", parent=styles["BodyText"], fontSize=10, leading=14, spaceAfter=6)

    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            story.append(Paragraph(stripped[2:], title_style))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], heading_style))
        elif stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], heading_style))
        elif stripped.startswith("---"):
            story.append(Spacer(1, 12))
        elif stripped.startswith("**") and stripped.endswith("**"):
            story.append(Paragraph(f"<b>{stripped[2:-2]}</b>", body_style))
        elif stripped.startswith("- "):
            story.append(Paragraph(f"\u2022 {stripped[2:]}", body_style))
        elif stripped:
            # Escape XML special characters for reportlab
            safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, body_style))

    doc.build(story)
    return filepath
